import os
import re
import pandas as pd
import streamlit as st
from docx import Document
import google.generativeai as genai
import traceback

# ---------------- Streamlit Setup ----------------
st.set_page_config(page_title="RnD DVT Test Planner", layout="wide")
st.title("RnD DVT Test Planner")

# ---------------- Repo Config ----------------
REPO_PATH = "."  # Current directory
DEFAULT_REQUIREMENTS_FILE = os.path.join(REPO_PATH, "dvt_requirements.xlsx")
HISTORY_DIR = os.path.join(REPO_PATH, "history")
os.makedirs(HISTORY_DIR, exist_ok=True)

# ---------------- API Key Loader ----------------
def load_api_key():
    if hasattr(st, "secrets"):
        api_key = st.secrets.get("google_gemini", {}).get("api_key")
        if api_key:
            return api_key
    api_key = os.environ.get("GOOGLE_API_KEY")
    if api_key:
        return api_key
    return None

api_key = load_api_key()
if api_key:
    genai.configure(api_key=api_key)
else:
    st.warning(
        "Google Gemini API key not found. AI suggestions will show 'AI suggestion failed'."
    )

# ---------------- File Readers ----------------
def read_requirements_file(file_path, uploaded=False):
    """
    Reads the requirements Excel file.
    - Only shows errors if uploaded=True (i.e., the user actually uploaded a file)
    """
    if uploaded and not file_path.endswith(".xlsx"):
        st.error("Please upload a valid Excel (.xlsx) file.")
        return None

    try:
        df = pd.read_excel(file_path)
        if df.shape[1] < 3:
            if uploaded:
                st.error("Excel file must have at least 3 columns: ID, Category, Description")
            return None
        return df
    except Exception as e:
        if uploaded:
            st.error(f"Failed to read Excel file: {e}")
        return None

def docx_to_text(file):
    doc = Document(file)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def load_rules_for_requirement(requirement_id):
    rule_file = os.path.join(REPO_PATH, f"{requirement_id}_Rule.docx")
    if os.path.exists(rule_file):
        try:
            doc = Document(rule_file)
            return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        except Exception as e:
            st.warning(f"Failed to read rule file {rule_file}: {e}")
            return []
    else:
        st.warning(f"No rules file found for {requirement_id}")
        return []

# ---------------- Token Normalization ----------------
def normalize_token(token):
    token = token.lower()
    match = re.match(r'(-?\d+\.?\d*)([a-z%]*)', token)
    if match:
        number_part = match.group(1)
        unit_part = match.group(2)
        normalized = f"{number_part}{unit_part}"
        normalized = re.sub(r'[^a-z0-9]', '', normalized)
        return normalized
    return re.sub(r'[^a-z0-9]', '', token)

def extract_check_items_robust(text):
    text = text.lower()
    number_matches = re.findall(r'-?\d+\.?\d*\s*[a-z%]*', text)
    alnum_matches = re.findall(r'\b[\w\-]{2,}\b', text)
    items = set()
    for n in number_matches + alnum_matches:
        cleaned = re.sub(r'[^a-z0-9]', '', n)
        if cleaned:
            items.add(cleaned)
    return items

def get_missing_rule_lines(rule_lines, plan_text):
    plan_tokens = extract_check_items_robust(plan_text)
    normalized_plan_tokens = {normalize_token(t) for t in plan_tokens}

    missing_lines = []
    for line in rule_lines:
        rule_tokens = re.findall(r'\b[\w\-\+\.]+\b', line)
        if any(normalize_token(token) not in normalized_plan_tokens for token in rule_tokens):
            missing_lines.append(line)
    return missing_lines

# ---------------- History ----------------
def load_history(requirement_id):
    if not os.path.exists(HISTORY_DIR):
        return ""
    history_files = [f for f in os.listdir(HISTORY_DIR) if f.startswith(requirement_id)]
    history_texts = []
    for file in history_files:
        with open(os.path.join(HISTORY_DIR, file), "r") as f:
            history_texts.append(f.read())
    return "\n".join(history_texts)

def save_history(requirement_id, missing_rule_lines, plan_text, ai_suggestions):
    timestamp = pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")
    history_file = os.path.join(HISTORY_DIR, f"{requirement_id}_{timestamp}.txt")
    with open(history_file, "w") as f:
        f.write("Rule Lines:\n" + "\n".join(missing_rule_lines) + "\n\n")
        f.write("Proposed Plan:\n" + plan_text + "\n\n")
        if isinstance(ai_suggestions, list):
            f.write("AI Suggestions:\n" + "\n".join(ai_suggestions))
        else:
            f.write("AI Suggestions:\n" + ai_suggestions)

# ---------------- Gemini AI Suggestions ----------------
def get_gemini_suggestions(plan_text, missing_rule_lines, requirement_id):
    if not missing_rule_lines:
        return "No missing rules; coverage complete"

    history = load_history(requirement_id)

    prompt = f"""
You are an engineering test coverage assistant, expert in RF and power systems.
Analyze the proposed test plan below and provide up to 5 concise suggestions to improve the test coverage.
Focus on critical areas such as safety, edge cases, performance metrics, and compliance with standards.
Use your expertise to suggest practical, implementable tests.
Include a short reasoning for each suggestion.

Review the past analyses below to build upon previous suggestions. You can repeat the previously suggested ideas if the test plan needs it. 

Past Analyses:
{history}

Proposed Test Plan:
{plan_text}

Missing Rules:
{chr(10).join(missing_rule_lines)}

Output ONLY the suggestions in the following exact markdown format, with no additional text or numbering:

- **Suggestion:** Concise suggestion description.
  - **Reasoning:** Short reasoning here.
"""

    try:
        if not api_key:
            return "AI suggestion failed: No API key configured"

        model = genai.GenerativeModel('gemini-2.5-pro')
        response = model.generate_content(prompt)
        ai_text = response.text
        if not ai_text.strip():
            return "AI suggestion failed"
        save_history(requirement_id, missing_rule_lines, plan_text, ai_text)
        return ai_text

    except Exception as e:
        tb = traceback.format_exc()
        return f"AI suggestion failed: {str(e)}\n\nTraceback: {tb}"

# ---------------- Step 1: Upload Requirements ----------------
uploaded_req_file = st.file_uploader(
    "Step 1: Upload Requirements Excel (.xlsx)",
    type=["xlsx"]
)

if uploaded_req_file:
    # Validate uploaded filename
    if uploaded_req_file.name.lower() != "dvt_requirements.xlsx":
        st.error("Please upload a valid requirement file named 'dvt_requirements.xlsx'.")
        st.stop()

    temp_req_path = os.path.join(REPO_PATH, "temp_requirements.xlsx")
    with open(temp_req_path, "wb") as f:
        f.write(uploaded_req_file.read())
    REQUIREMENTS_FILE = temp_req_path
    df = read_requirements_file(REQUIREMENTS_FILE, uploaded=True)
else:
    # Only use default if it exists
    if os.path.exists(DEFAULT_REQUIREMENTS_FILE):
        REQUIREMENTS_FILE = DEFAULT_REQUIREMENTS_FILE
        df = read_requirements_file(REQUIREMENTS_FILE, uploaded=False)
    else:
        st.warning("No requirements file uploaded and default file not found.")
        st.stop()

if df is None:
    st.stop()

# ---------------- Step 2: Requirement ID ----------------
requirement_ids = df.iloc[:, 0].astype(str).tolist()
categories = df.iloc[:, 1].astype(str).tolist()
descriptions = df.iloc[:, 2].astype(str).tolist()
id_to_category = {rid.upper(): cat for rid, cat in zip(requirement_ids, categories)}
id_to_description = {rid.upper(): desc for rid, desc in zip(requirement_ids, descriptions)}

user_input = st.text_input("Step 2: Enter the Requirement ID (e.g. DS-1)").strip().upper()
valid_id = False
if user_input:
    if user_input in id_to_description:
        valid_id = True
        st.success(
            f"**{user_input}**\n\n"
            f"**Category:** {id_to_category[user_input]}\n\n"
            f"**Description:** {id_to_description[user_input]}"
        )
    else:
        st.error("No match found for that Requirement ID.")

# ---------------- Step 3: Upload Test Plan ----------------
if valid_id:
    uploaded_plan_file = st.file_uploader(
        "Step 3: Upload Proposed Test Plan (.docx or .txt)", 
        type=["docx", "txt"]
    )

    if uploaded_plan_file:
        # Check filename corresponds to requirement ID
        if user_input not in uploaded_plan_file.name:
            st.error("Please upload a valid test plan with requirement ID in filename.")
        else:
            if uploaded_plan_file.name.endswith(".docx"):
                plan_lines = docx_to_text(uploaded_plan_file)
            else:
                try:
                    plan_lines = uploaded_plan_file.read().decode("utf-8").splitlines()
                except UnicodeDecodeError:
                    plan_lines = uploaded_plan_file.read().decode("latin1").splitlines()

            plan_text = "\n".join(plan_lines)

            if st.button("Analyze Test Plan"):
                with st.spinner("Analyzing test plan..."):
                    rule_lines = load_rules_for_requirement(user_input)
                    missing_lines = get_missing_rule_lines(rule_lines, plan_text) if rule_lines else []

                    st.markdown("## Your Proposed Test Plan")
                    for line in plan_lines:
                        if line.strip():
                            st.markdown(f"- {line.strip()}")

                    st.markdown("## Based on Past Test Cases")
                    if missing_lines:
                        for line in missing_lines:
                            st.markdown(f"- {line}")
                    else:
                        st.success("All rule lines are fully covered in the proposed plan!")

                    ai_suggestions = get_gemini_suggestions(plan_text, missing_lines, user_input)
                    st.markdown("## AI Suggestions")
                    st.markdown(ai_suggestions)
